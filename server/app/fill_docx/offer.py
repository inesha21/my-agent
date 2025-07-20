from docx import Document
from typing import Dict
import io

def replace_text_in_paragraph(paragraph, data: Dict[str, str]):
    full_text = ''.join(run.text for run in paragraph.runs)
    replaced = False
    for key, value in data.items():
        placeholder = f"[{key}]"
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, value)
            replaced = True
    if replaced:
        for run in paragraph.runs:
            run.text = ''
        paragraph.runs[0].text = full_text

def generate_offer_letter_docx(fields: Dict[str, str], template_path: str = "docx/templates/offer.docx") -> bytes:
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, fields)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, fields)

    # Save to in memory bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def generate_confirmation_letter_docx(fields: Dict[str, str], template_path: str = "docx/templates/confirmation.docx") -> bytes:
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, fields)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, fields)

    # Save to in memory bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
